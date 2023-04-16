import re
import docx


def extract_theory(file_path: str) -> list[str]:
    doc = docx.Document(file_path)

    list_number_elements = []
    count = 0
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'List Number' or paragraph.style.name == 'List Paragraph':
            count += 1
            list_number_elements.append(paragraph.text)

    return list_number_elements


def extract_practice(file_path: str) -> list[str]:
    doc = docx.Document(file_path)
    task_elements = []

    pattern = re.compile(r'^Задание\s*\d*', re.IGNORECASE)
    pattern_english = re.compile(r'^Task\s*\d*', re.IGNORECASE)

    for paragraph in doc.paragraphs:
        if pattern.match(paragraph.text):
            text_without_task_word = paragraph.text.split(maxsplit=1)[1]
            task_elements.append('Практическое задание ' + text_without_task_word)
        elif pattern_english.match(paragraph.text):
            text_without_task_word = paragraph.text.split(maxsplit=1)[1]
            task_elements.append('Practical task ' + text_without_task_word)

    return task_elements
