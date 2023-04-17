from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from app.models.ticket import Ticket


def generate_docx(file_path: str, tickets: list[Ticket]):
    document = Document()

    for ticket in tickets:
        p = document.add_paragraph()
        break_run = p.add_run()
        break_run.add_break()
        r = p.add_run(ticket.ticket_name)
        font = r.font
        font.name = 'Times New Roman'
        font.size = Pt(15)
        r.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        tasks = ticket.theory + ticket.practice
        for index, task in enumerate(tasks):
            p = document.add_paragraph()
            r = p.add_run(f'{index + 1}.  {task.description}')
            font = r.font
            font.name = 'Times New Roman'
            font.size = Pt(16)
            r.bold = False

    document.save(file_path)
